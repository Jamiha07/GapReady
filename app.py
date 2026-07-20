"""
Step 3: Flask backend.
Wraps our working resume-extraction + Groq prompt logic into a web endpoint
that a frontend (or a tool like Postman) can actually call.

FIXES APPLIED IN THIS VERSION:
1. API key moved to an environment variable (GROQ_API_KEY) — never hardcode it again.
2. Resume is read directly from memory — no temp file on disk (removes the unsafe-filename risk).
3. JD validation hardened + gap-grounding rule added to the system prompt.
4. Groq JSON mode enabled + JSON is parsed server-side, so the frontend always gets clean data.
5. /sample-answer wrapped in try/except.
6. Server-side .docx extension check + 10MB upload limit actually enforced.
7. temperature set for more consistent structured output; truncation is now flagged to the model.
"""

import json
import os
import re

from flask import Flask, request, jsonify
from flask_cors import CORS
from groq import Groq
from docx import Document
from nlp_analysis import analyze, cluster_gaps

# FIX 1: read the key from an environment variable instead of hardcoding it.
# Before running:  set GROQ_API_KEY=your_new_key   (Windows CMD)
#             or:  $env:GROQ_API_KEY="your_new_key" (PowerShell)
API_KEY = os.environ.get("Interview_qs_groq_key")
if not API_KEY:
    raise RuntimeError(
        "interview_qs_groq_key environment variable is not set. "
        "Set it before running app.py — do not paste the key into the code."
    )

app = Flask(__name__, static_folder="static")
CORS(app)


@app.route("/")
def home():
    """Serve the frontend page — the whole app is one deployment."""
    return app.send_static_file("index.html")


# FIX 6: enforce the 10MB limit the UI advertises. Requests larger than this
# are rejected by Flask automatically with a 413 error.
app.config["MAX_CONTENT_LENGTH"] = 10 * 1024 * 1024

client = Groq(api_key=API_KEY)

SYSTEM_PROMPT = """
You are an AI interview preparation assistant for students preparing for internships and entry-level roles. You will be given two inputs: (1) a student's resume text, and (2) a job/internship description they are applying for.

You MUST follow this process in exact order:

STEP 1 — INPUT VALIDATION (do this first, always):
Check the RESUME input specifically for resume-like structure and breadth. A genuine resume covers MULTIPLE distinct areas of a person's background BRIEFLY — it is not deep, extensive documentation of a single project or topic. Specifically, a valid resume must show clear evidence of AT LEAST THREE of the following, each appearing as its own distinct, relatively brief section (not one topic explored at length):
(a) contact information (email, phone, LinkedIn, etc.) — a name alone is NOT sufficient, since names appear on many document types (reports, papers, assignments);
(b) an education section listing institution(s) and degree/program information;
(c) a skills list (technologies, tools, or competencies, typically as short items, not prose);
(d) a work experience or internship section with dates and role titles;
(e) a projects section that briefly lists MULTIPLE projects (not one single project explored in deep technical detail across many paragraphs or pages).

If the document is a university course submission, project report, thesis, research paper, or any single-project deep-dive — even if it includes the student's name, course details, or a CMS/roll number, and even if it describes a project the student genuinely built — this does NOT count as a resume unless it ALSO independently shows breadth across contact info, education, skills, and/or multiple distinct experiences. A single project described at length, with a name and course header attached, is still not a resume.

Separately, validate the JOB DESCRIPTION input just as strictly. A genuine job/internship posting must contain AT LEAST TWO of the following:
(a) a named role or position title (e.g., "Full Stack Intern", "ML Engineer");
(b) a responsibilities/duties section describing what the hire will do;
(c) a requirements/qualifications section describing what the employer is looking for;
(d) company, location, duration, or compensation details.
Text that is merely TECHNICAL in nature does NOT qualify — code review feedback, programming advice, documentation, tutorials, articles, chat or email excerpts, and numbered lists of suggestions are NOT job descriptions, even if they mention tools, languages, or frameworks. The test is: does this text read as an EMPLOYER describing a role they are hiring for? If not, it fails, and you must set "valid_input": false with an error_message explaining that the second input does not appear to be a job description (e.g., "The pasted text appears to be technical notes or advice, not a job posting — it has no role title, responsibilities, or requirements section. Please paste an actual job or internship description.").

If EITHER the resume fails its structural/breadth check above OR the job description fails its check, stop here and output only the JSON with "valid_input": false and a clear error_message explaining specifically which input was the problem and why. Do not proceed to Step 2 or Step 3. Never fabricate or assume a person's project/work experience from a document that is not actually their resume, and never invent job requirements from text that is not actually a job description.

CRITICAL — MISMATCH IS NOT INVALIDITY: Validation checks each input's STRUCTURE independently. It does NOT compare the resume and job description to each other. A structurally valid resume + a structurally valid job description from a completely different field (e.g. a computer science resume with a finance, marketing, or nursing internship posting) is a VALID input pair — students legitimately apply across fields, and analyzing a poor fit is exactly this tool's purpose. In such cases you MUST proceed to Step 2 and beyond; the mismatch will naturally appear as a low match and many gaps. Never set "valid_input": false because the two inputs belong to different fields, and never ask the user to make the job description "match" the resume.

STEP 2 — GAP VERIFICATION (mandatory, done explicitly, before any questions are written):
You will be given a CANDIDATE GAPS list — terms from the job description that a statistical word-matching analysis could not find in the resume. This list is intentionally over-inclusive: it contains real gaps, but also noise (sentence fragments like "version" or "control" that are not skills) and false positives (skills the resume DOES demonstrate but under different wording — e.g. the candidate list may contain "backend" while the resume shows Flask REST API work, which IS backend development).

Your job is to VERIFY each candidate:
1. Discard items that are not real skills or requirements (fragments, generic words).
2. Discard items the resume clearly demonstrates under different wording — re-read the resume carefully for semantic evidence before keeping any item.
3. Merge related fragments into one properly-named gap where appropriate (e.g. "docker" + "containerization" become one gap: "Containerization with Docker").
4. Keep only genuine gaps: requirements explicitly stated in the job description text with no evidence in the resume.

You must also report your classifications — this is used to compute an accurate match score, so be precise:
- Every candidate term you discarded because the resume DOES demonstrate it under different wording (rule 2) goes into a "dismissed_candidates" list, copied as the term appeared in the candidate list.
- Every candidate term you discarded because it is NOT a real skill or requirement at all (rule 1 — sentence fragments, generic words like "startup" or "practical") goes into a "noise_terms" list, copied as the term appeared in the candidate list.
Every candidate term must end up in exactly one place: dismissed_candidates (covered), noise_terms (not a requirement), or reflected in identified_gaps (genuine gap). Do not silently drop any candidate. Never put the same term in more than one list. dismissed_candidates is ONLY for genuine skills/requirements that the resume clearly demonstrates — generic words, fragments, and filler (e.g. "startup", "site", "practical", "user", "improve") are ALWAYS noise_terms, never dismissed_candidates. When unsure whether a term is a covered skill or noise, classify it as noise — never award coverage without clear resume evidence. Candidate terms that correspond to your identified_gaps (even after rewording or merging) must NOT appear in noise_terms or dismissed_candidates — they are already accounted for as gaps.

Every gap in your final identified_gaps list MUST correspond to a requirement explicitly stated in the job description text. Never invent gaps from general knowledge of what similar roles usually require. If, after verification, no genuine gaps remain, output an empty identified_gaps list and skip gap questions (3a) entirely — do not manufacture gaps to fill space.

You must ALSO produce a "covered_requirements" list: the distinct skills/requirements explicitly stated in the job description that the resume DOES clearly demonstrate (under any wording). Phrase each the same way you phrase gaps — short named requirements, not raw words. Together, covered_requirements + identified_gaps must account for every distinct real requirement in the job description: each requirement appears in exactly one of the two lists. The match score shown to the student is computed directly as covered ÷ (covered + gaps), so missing or duplicated requirements will distort their score — be complete and precise.

STEP 3 — QUESTION GENERATION (only after Steps 1 and 2 are complete):
Generate questions in this exact order:

3a. GAP QUESTIONS FIRST (non-negotiable): For EACH item in your Step 2 gap list, generate exactly one question directly testing that specific gap. You must produce one question per gap. Tag each with "based_on": "gap".

3b. PROJECT QUESTIONS SECOND: After all gap questions are written, generate 2-3 additional questions referencing the student's actual named projects from their resume. Tag these with "based_on": "project".

3c. APPLY THESE RULES TO ALL QUESTIONS ABOVE (both 3a and 3b):
- Difficulty must match the job description's actual requirements, not the student's apparent seniority — never soften a question out of consideration for inexperience.
- CATEGORY CLASSIFICATION: Classify each question strictly by its actual content, not by which project/gap it's based on. A question is "technical" if it asks about implementation, tools, architecture, code, or technical trade-offs — even if it references a project by name. A question is "behavioral" ONLY if it asks about teamwork, communication, conflict, motivation, or soft-skill situations. Do not default to "behavioral" just because a question mentions a personal project.
- Ask ONE topic per question — never combine two unrelated requirements into a single confusing question.
- Include one natural follow-up question per main question.
- Include a short hint (2-3 bullet points, NOT a full answer) per question.

3d. MANDATORY BEHAVIORAL QUESTIONS: In addition to the gap and project questions above, generate at least 2 behavioral questions (teamwork, communication, conflict, motivation, or soft-skill situations) — even if no natural behavioral gap exists between the resume and JD. Tag these with "based_on": "general".

STEP 4 — SELF-CHECK (mandatory, before you output anything):
1. Count the number of items in your identified_gaps list. Count the number of questions tagged "based_on": "gap". These two numbers MUST match — add any missing gap question before finalizing.
2. Count the number of questions with "category": "behavioral". This MUST be at least 2. If fewer than 2 exist, add additional behavioral questions before finalizing output.
3. Do not output a response that fails either check above.

STEP 5 — FORMAT CHECK (mandatory, before you output anything):
Verify every "hint" field is a single plain string (e.g., "Mention X, Y, and Z"), NOT a list or array of separate bullet points. If a hint currently exists as a list, rewrite it as one combined string with items separated by commas or semicolons before finalizing output. Every field in the schema must match its specified type exactly — strings must be strings, not arrays.

OUTPUT FORMAT: Respond ONLY in the following JSON structure, with no extra commentary before or after:

{
  "valid_input": true or false,
  "error_message": "string, only present if valid_input is false",
  "identified_gaps": ["string", "string", ...],
  "covered_requirements": ["string", ...],
  "dismissed_candidates": ["string", ...],
  "noise_terms": ["string", ...],
  "questions": [
    {
      "category": "technical" or "behavioral",
      "based_on": "gap" or "project" or "general",
      "question": "string",
      "hint": "string",
      "followup": "string"
    }
  ]
}
"""


def extract_resume_text(file_obj):
    """Reads paragraphs and tables from a .docx file.
    FIX 2: python-docx accepts a file-like object directly, so we pass the
    uploaded file straight in — no saving to disk, no temp folder, no cleanup."""
    doc = Document(file_obj)
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    full_text.append(cell.text)
    return "\n".join(full_text)


@app.route("/generate-questions", methods=["POST"])
def generate_questions():
    # 1. Check a resume file was actually sent
    if "resume" not in request.files:
        return jsonify({"error": "No resume file uploaded"}), 400

    resume_file = request.files["resume"]
    job_description = request.form.get("job_description", "").strip()

    if not job_description:
        return jsonify({"error": "No job description provided"}), 400

    # FIX 6: enforce .docx on the server, not just in the browser.
    if not resume_file.filename.lower().endswith(".docx"):
        return jsonify({"error": "Please upload a .docx file — other formats are not supported."}), 400

    # 2. Extract resume text directly from the uploaded file (FIX 2: no temp file).
    try:
        resume_text = extract_resume_text(resume_file)
    except Exception as e:
        return jsonify({"error": f"Could not read resume file: {str(e)}"}), 400

    if not resume_text.strip():
        return jsonify({"error": "The uploaded file contains no readable text."}), 400

    # 2b. Truncate resume text if it's extremely long (e.g. someone uploaded a research
    # paper instead of a resume) — avoids hitting Groq's per-minute token limit.
    # FIX 7: the truncation is now flagged so the model knows the text was cut.
    MAX_RESUME_CHARS = 6000
    if len(resume_text) > MAX_RESUME_CHARS:
        resume_text = resume_text[:MAX_RESUME_CHARS] + "\n[NOTE: text truncated here — original document was much longer]"

    # 3. NLP layer: compute coverage score, similarity, and candidate gaps.
    # This is pure local math (no API call, costs nothing) and runs BEFORE Groq.
    nlp = analyze(resume_text, job_description)

    # 4. Build the user message and call Groq — now with a third block:
    # the candidate gap list for the LLM to verify (see Step 2 of the prompt).
    user_message = f"""
RESUME:
{resume_text}

JOB DESCRIPTION:
{job_description}

CANDIDATE GAPS (from statistical analysis — verify each one against the resume):
{nlp["candidate_gaps"]}
"""

    try:
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": user_message}
            ],
            # FIX 7: lower temperature = more consistent, rule-following output
            temperature=0.4,
            # FIX 4: JSON mode — the model is forced to return valid JSON
            response_format={"type": "json_object"},
        )
    except Exception as e:
        return jsonify({"error": f"The request could not be processed (it may be too large, or the free API limit was hit). Details: {str(e)}"}), 500

    # 4. FIX 4: parse the JSON here on the server, so the frontend never has to
    # deal with raw model text. If parsing somehow still fails, return a clean error.
    raw = response.choices[0].message.content
    try:
        parsed = json.loads(raw)
    except (json.JSONDecodeError, TypeError):
        return jsonify({"error": "The AI returned an unreadable response. Please try again."}), 502

    # 4b. ENFORCE the one-question-per-gap rule in code (the prompt asks for it,
    # but LLMs occasionally ignore their own self-check). Instead of calling
    # Groq a second time (costs double tokens), we just trim silently so the
    # counts always match what the user sees.
    def gap_question_count(p):
        return sum(1 for q in p.get("questions", []) if q.get("based_on") == "gap")

    if parsed.get("valid_input") and len(parsed.get("identified_gaps", [])) != gap_question_count(parsed):
        print(f"DEBUG mismatch: {len(parsed.get('identified_gaps', []))} gaps vs "
              f"{gap_question_count(parsed)} gap questions — trimming (no retry, saves tokens)")
        parsed["identified_gaps"] = parsed.get("identified_gaps", [])[:gap_question_count(parsed)]

    # 5. MATCH SCORE — requirement-level (final design).
    # score = covered requirements / (covered + missing requirements)
    # This makes the score and the gaps panel two views of the SAME lists,
    # so they can never contradict each other again.
    n_covered = len(parsed.get("covered_requirements", []) or [])
    n_gaps = len(parsed.get("identified_gaps", []) or [])

    if n_covered + n_gaps > 0:
        verified_score = round((n_covered / (n_covered + n_gaps)) * 100, 1)
    else:
        # Fallback (LLM sent no requirement lists): word-level verified score,
        # same math as the previous version.
        word_matched = len(nlp["covered_terms"])
        total_terms = word_matched + len(nlp["candidate_gaps"])

        def words_of(items):
            text = " ".join(str(t).lower() for t in (items or []))
            return set(re.findall(r"[a-z0-9\+\#\.]+", text))

        dismissed_words = words_of(parsed.get("dismissed_candidates"))
        noise_words = words_of(parsed.get("noise_terms"))
        gap_words = words_of(parsed.get("identified_gaps"))

        dismissed = [t for t in nlp["candidate_gaps"]
                     if t.lower() in dismissed_words and t.lower() not in gap_words]
        noise = [t for t in nlp["candidate_gaps"]
                 if t.lower() in noise_words
                 and t.lower() not in dismissed_words
                 and t.lower() not in gap_words]

        effective_total = total_terms - len(noise)
        if effective_total > 0:
            verified_score = round(((word_matched + len(dismissed)) / effective_total) * 100, 1)
            verified_score = min(verified_score, 100.0)
        else:
            verified_score = 0

    print(f"DEBUG score: covered_reqs={n_covered} gaps={n_gaps} score={verified_score}")

    # 6. K-Means: group the final gaps into related categories for display.
    # Local math only, no API cost. Returns None for short lists (frontend
    # then shows the flat list exactly as before).
    gap_groups = cluster_gaps(parsed.get("identified_gaps", []))

    return jsonify({
        "result": parsed,
        "gap_groups": gap_groups,                   # K-Means clusters (or null)
        "coverage_score": verified_score,           # headline: word matches + semantic matches
        "keyword_score": nlp["coverage_score"],     # small print: exact-word matches only
        "similarity": nlp["similarity"],            # small print: overall text overlap
    })


@app.route("/chat", methods=["POST"])
def chat():
    """Internship support chatbot. Receives the conversation history AND an
    optional summary of the user's latest GapReady analysis (score, gaps,
    covered skills) so it can give specific advice instead of generic tips."""
    data = request.get_json(silent=True) or {}
    history = data.get("messages", [])
    context = data.get("context") or {}

    # Basic shape check: history must be a list of {role, content} dicts.
    if not isinstance(history, list) or not history:
        return jsonify({"error": "No message provided"}), 400

    # Keep only the fields we expect, cap history length to control tokens.
    clean_history = []
    for m in history[-12:]:                     # last 12 turns is plenty of context
        role = m.get("role")
        content = str(m.get("content", ""))[:2000]   # cap each message's size
        if role in ("user", "assistant") and content.strip():
            clean_history.append({"role": role, "content": content})

    # Build a short, safe summary of the user's latest analysis — NOT the
    # full resume text (privacy + token cost). Only score/gaps/covered.
    context_block = ""
    if context.get("coverage_score") is not None:
        gaps = context.get("gaps", [])[:15]       # cap length, keep tokens low
        covered = context.get("covered", [])[:15]
        context_block = f"""

THE STUDENT'S LATEST GAPREADY RESULTS (use this if their question relates to it):
Match score: {context.get('coverage_score')}%
Skills the resume covers for this role: {covered}
Skills missing for this role: {gaps}
"""

    chat_system_prompt = """
You are GapBot, GapReady's built-in support assistant. GapReady is a web tool where
students upload their resume and a job description, and receive: a match score, a
list of skill gaps, and tailored interview questions with hints and sample answers.

Your job is to help students with:
- Interview preparation advice (how to answer question types, STAR method, common mistakes)
- Understanding and acting on their GapReady results (what a gap means, how to close one,
  how to talk about a missing skill honestly in an interview)
- Internship application guidance (resumes, cover letters, follow-up etiquette)
- Explaining technical terms that appear in job descriptions or interview questions
- How to use the GapReady tool itself

If a "LATEST GAPREADY RESULTS" block is provided below, you may reference the student's
actual score/gaps/covered skills by name to give specific advice. You do NOT have their
full resume — only this summary. Never invent details beyond what's given.

FORMATTING RULES (important):
- Prefer short bullet points or a few short headed sections over one long paragraph.
- Keep language simple and direct — this is a first-year student, not a formal essay.
- Only use a plain short paragraph for very simple one-line answers; use structure
  (bullets/headings) whenever the answer has more than one part or step.
- Never pad with filler sentences. Every line should say something useful.

Rules:
- Stay on topic. If asked something unrelated to internships, careers, interviews, or
  GapReady (e.g. write my homework, general coding tasks, politics), politely decline in
  one sentence and steer back to interview/internship help.
- Be encouraging but honest — no empty hype.
""" + context_block

    try:
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "system", "content": chat_system_prompt}] + clean_history,
            temperature=0.6,
            max_tokens=380,           # kept modest to protect the daily token budget
        )
    except Exception as e:
        return jsonify({"error": f"The chat service is unavailable right now. Details: {str(e)}"}), 500

    return jsonify({"reply": response.choices[0].message.content})


@app.route("/sample-answer", methods=["POST"])
def sample_answer():
    data = request.get_json(silent=True) or {}
    question = data.get("question", "").strip()
    hint = data.get("hint", "").strip()

    if not question:
        return jsonify({"error": "No question provided"}), 400

    prompt = f"""
You are helping a student prepare for an internship interview. Given this interview question and hint, write ONE realistic, well-structured sample answer a strong candidate might give — written in first person, natural spoken tone, 3-5 sentences. Do not add commentary before or after the answer, just the answer itself.

Question: {question}
What a strong answer should cover: {hint}
"""

    # FIX 5: wrapped in try/except so a Groq failure returns a clean error
    # instead of crashing with an unhandled 500.
    try:
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}]
        )
    except Exception as e:
        return jsonify({"error": f"Could not generate a sample answer right now. Details: {str(e)}"}), 500

    return jsonify({"answer": response.choices[0].message.content})


if __name__ == "__main__":
    # This block only runs locally (python app.py).
    # On Render, gunicorn runs the app instead and ignores this.
    app.run(debug=True, port=5000)